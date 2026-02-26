import logging
import asyncio
from typing import List
from mcp.server.fastmcp import FastMCP
from exchangelib import Message, Mailbox, Q, FolderCollection
from exchangelib.items import Item
from exchangelib.errors import ErrorItemNotFound
import bs4

from .client import get_ews_client
from .utils import build_email_body
from .idempotency import IdempotencyManager

logger = logging.getLogger("ews_mcp")

mcp = FastMCP("email-exchange-mcp")
idempotency = IdempotencyManager()

def html_to_text(html: str) -> str:
    """Safely convert HTML string to text."""
    if not html:
        return ""
    try:
        soup = bs4.BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n").strip()
    except Exception:
        return html

def get_folder_by_name(account, folder_name: str):
    """Resolve well-known folder names or search by name."""
    name_lower = folder_name.lower().strip()
    if name_lower in ["inbox", "收件箱"]: return account.inbox
    if name_lower in ["sent", "sentitems", "已发送"]: return account.sent
    if name_lower in ["drafts", "草稿箱"]: return account.drafts
    if name_lower in ["deleteditems", "已删除"]: return account.trash
    if name_lower in ["junk", "垃圾邮件"]: return account.junk
    
    # Custom folder fallback
    for f in account.root.walk():
        if f.name and f.name.lower() == name_lower:
            return f
    raise ValueError(f"Folder '{folder_name}' not found.")

def _format_item(item, fetch_body=False):
    """Serialize exchangelib item to dict."""
    res = {
        "id": item.id,
        "subject": item.subject or "(No Subject)",
        "sender": item.sender.email_address if hasattr(item, 'sender') and item.sender else "Unknown",
        "datetime_received": item.datetime_received.isoformat() if item.datetime_received else None,
        "is_read": item.is_read if hasattr(item, 'is_read') else True,
        "has_attachments": item.has_attachments if hasattr(item, 'has_attachments') else False
    }
    if fetch_body:
        res["body"] = html_to_text(item.body) if item.body else ""
        res["html_body"] = str(item.body) if item.body else ""
    return res

# ---------------------------------------------------------
# Read Operations
# ---------------------------------------------------------

@mcp.tool()
def list_messages(folder_name: str = "inbox", limit: int = 20, fetch_body: bool = False) -> str:
    """List newest messages in a folder (e.g. 'inbox', 'sent')."""
    account = get_ews_client()
    folder = get_folder_by_name(account, folder_name)
    
    qs = folder.all().order_by('-datetime_received')
    
    if fetch_body:
        qs = qs.only('id', 'subject', 'sender', 'datetime_received', 'is_read', 'has_attachments', 'body')
    else:
        qs = qs.only('id', 'subject', 'sender', 'datetime_received', 'is_read', 'has_attachments')
    
    messages = []
    for item in qs[:limit]:
        messages.append(_format_item(item, fetch_body=fetch_body))
        
    import json
    return json.dumps({"folder": folder.name, "count": len(messages), "messages": messages}, ensure_ascii=False)


@mcp.tool()
def get_message_details(message_id: str) -> str:
    """Get full details of an email by ID."""
    account = get_ews_client()
    try:
        item = account.root.get(id=message_id)
        res = _format_item(item, fetch_body=True)
        if isinstance(item, Message):
            res["to_recipients"] = [r.email_address for r in item.to_recipients] if item.to_recipients else []
            res["cc_recipients"] = [r.email_address for r in item.cc_recipients] if item.cc_recipients else []
            res["datetime_sent"] = item.datetime_sent.isoformat() if item.datetime_sent else None
        
        import json
        return json.dumps({"success": True, "message": res}, ensure_ascii=False)
    except ErrorItemNotFound:
        return f'{{"success": false, "error": "Message ID {message_id} not found."}}'


@mcp.tool()
def search_messages(query: str, folder_name: str = "inbox", limit: int = 10, fetch_body: bool = False) -> str:
    """Search messages using keywords (e.g. 'subject:Project')."""
    account = get_ews_client()
    folder = get_folder_by_name(account, folder_name)
    
    # Exchangelib natively supports AQS by just passing string to filter/all
    qs = folder.filter(query).order_by('-datetime_received')
    
    if fetch_body:
        qs = qs.only('id', 'subject', 'sender', 'datetime_received', 'is_read', 'has_attachments', 'body')
        
    messages = [_format_item(item, fetch_body) for item in qs[:limit]]
    import json
    return json.dumps({"success": True, "query": query, "count": len(messages), "messages": messages}, ensure_ascii=False)


@mcp.tool()
def get_conversation_thread(message_id: str, limit: int = 20) -> str:
    """Get all messages in the same conversation thread as the given message."""
    account = get_ews_client()
    try:
        seed_item = account.root.get(id=message_id)
        if not hasattr(seed_item, 'conversation_id') or not seed_item.conversation_id:
            import json
            return json.dumps({"success": False, "error": "No conversation thread found for this message."}, ensure_ascii=False)
        
        conv_id = seed_item.conversation_id.id
        qs = account.root.filter(conversation_id=conv_id).order_by('-datetime_received')[:limit]
        
        messages = [_format_item(item, fetch_body=False) for item in qs]
        import json
        return json.dumps({"success": True, "conversation_id": conv_id, "count": len(messages), "messages": messages}, ensure_ascii=False)
    except Exception as e:
        import json
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)


@mcp.tool()
def list_attachments(message_id: str) -> str:
    """List information about all attachments of an email."""
    account = get_ews_client()
    try:
        item = account.root.get(id=message_id)
        attachments = []
        for att in item.attachments:
            from exchangelib import FileAttachment
            if isinstance(att, FileAttachment):
                attachments.append({
                    "name": att.name,
                    "size": att.size,
                    "content_type": att.content_type,
                    "is_inline": att.is_inline
                })
        import json
        return json.dumps({"success": True, "message_id": message_id, "attachments": attachments}, ensure_ascii=False)
    except Exception as e:
        import json
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)


@mcp.tool()
def get_attachment_content(message_id: str, attachment_name: str) -> str:
    """Extract text content from an attachment (supports txt, csv, html, json, md, pdf, docx, xlsx)."""
    account = get_ews_client()
    try:
        item = account.root.get(id=message_id)
        import json
        for att in item.attachments:
            if att.name == attachment_name:
                content = att.content  # Binary content
                ext = att.name.split('.')[-1].lower() if '.' in att.name else ''
                
                if ext in ['txt', 'csv', 'log', 'json', 'md']:
                    return json.dumps({"success": True, "name": att.name, "content": content.decode('utf-8', errors='ignore')}, ensure_ascii=False)
                elif ext == 'html':
                    text = html_to_text(content.decode('utf-8', errors='ignore'))
                    return json.dumps({"success": True, "name": att.name, "content": text}, ensure_ascii=False)
                elif ext == 'pdf':
                    try:
                        import io
                        from pypdf import PdfReader
                        f = io.BytesIO(content)
                        reader = PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return json.dumps({"success": True, "name": att.name, "content": text.strip()}, ensure_ascii=False)
                    except Exception as pdf_err:
                        return json.dumps({"success": False, "error": f"Failed to parse PDF: {str(pdf_err)}"}, ensure_ascii=False)
                elif ext == 'docx':
                    try:
                        import io
                        import docx
                        f = io.BytesIO(content)
                        doc = docx.Document(f)
                        text = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text.append(para.text)
                        for table in doc.tables:
                            for row in table.rows:
                                text.append(" | ".join([cell.text.strip() for cell in row.cells]))
                        return json.dumps({"success": True, "name": att.name, "content": "\n".join(text)}, ensure_ascii=False)
                    except Exception as docx_err:
                        return json.dumps({"success": False, "error": f"Failed to parse DOCX: {str(docx_err)}"}, ensure_ascii=False)
                elif ext == 'xlsx':
                    try:
                        import io
                        from openpyxl import load_workbook
                        f = io.BytesIO(content)
                        wb = load_workbook(filename=f, data_only=True)
                        text = []
                        for sheetname in wb.sheetnames:
                            sheet = wb[sheetname]
                            text.append(f"--- Sheet: {sheetname} ---")
                            for row in sheet.iter_rows(values_only=True):
                                # Filter out completely empty rows to save space
                                if any(cell is not None for cell in row):
                                    text.append(" | ".join([str(cell) if cell is not None else "" for cell in row]))
                        return json.dumps({"success": True, "name": att.name, "content": "\n".join(text)}, ensure_ascii=False)
                    except Exception as xlsx_err:
                        return json.dumps({"success": False, "error": f"Failed to parse XLSX: {str(xlsx_err)}"}, ensure_ascii=False)
                else:
                    return json.dumps({"success": False, "error": f"Unsupported file type: {ext}. Supported: txt, csv, html, json, md, pdf, docx, xlsx."}, ensure_ascii=False)
        
        return json.dumps({"success": False, "error": f"Attachment '{attachment_name}' not found."}, ensure_ascii=False)
    except Exception as e:
        import json
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)


# ---------------------------------------------------------
# Write Operations
# ---------------------------------------------------------

@mcp.tool()
def send_email(
    to_recipients: str, 
    subject: str, 
    body: str, 
    idempotency_key: str, 
    cc_recipients: str = "", 
    use_signature: bool = True
) -> str:
    """Send an email. to/cc recipients should be comma-separated strings."""
    idempotency.lock(idempotency_key)
    try:
        account = get_ews_client()
        html_body = build_email_body(body, use_signature)
        
        to_list = [r.strip() for r in to_recipients.split(",") if r.strip()]
        cc_list = [r.strip() for r in cc_recipients.split(",") if r.strip()]
        
        msg = Message(
            account=account,
            folder=account.sent,
            subject=subject,
            body=html_body,
            to_recipients=[Mailbox(email_address=addr) for addr in to_list],
            cc_recipients=[Mailbox(email_address=addr) for addr in cc_list]
        )
        msg.send_and_save()
        
        idempotency.mark_success(idempotency_key)
        import json
        return json.dumps({"success": True, "action": "EmailSent", "idempotencyKey": idempotency_key})
    except Exception as e:
        idempotency.mark_failed(idempotency_key)
        logger.error(f"Failed to send email: {e}")
        raise


@mcp.tool()
def save_draft(
    to_recipients: str, 
    subject: str, 
    body: str, 
    idempotency_key: str, 
    cc_recipients: str = "", 
    use_signature: bool = True
) -> str:
    """Save a draft email. to/cc recipients should be comma-separated strings."""
    idempotency.lock(idempotency_key)
    try:
        account = get_ews_client()
        html_body = build_email_body(body, use_signature)
        
        to_list = [r.strip() for r in to_recipients.split(",") if r.strip()]
        cc_list = [r.strip() for r in cc_recipients.split(",") if r.strip()]
        
        msg = Message(
            account=account,
            folder=account.drafts,
            subject=subject,
            body=html_body,
            to_recipients=[Mailbox(email_address=addr) for addr in to_list],
            cc_recipients=[Mailbox(email_address=addr) for addr in cc_list]
        )
        msg.save()
        
        idempotency.mark_success(idempotency_key)
        import json
        return json.dumps({"success": True, "action": "DraftSaved", "idempotencyKey": idempotency_key, "messageId": msg.id})
    except Exception as e:
        idempotency.mark_failed(idempotency_key)
        logger.error(f"Failed to save draft: {e}")
        raise


@mcp.tool()
def reply_email(
    message_id: str, 
    body: str, 
    reply_all: bool, 
    idempotency_key: str, 
    use_signature: bool = True
) -> str:
    """Reply to an email by ID."""
    idempotency.lock(idempotency_key)
    try:
        account = get_ews_client()
        original_msg = account.root.get(id=message_id) # fetches the full message
        html_body = build_email_body(body, use_signature)
        
        if reply_all:
            original_msg.reply_all(subject=f"Re: {original_msg.subject}", body=html_body)
        else:
            original_msg.reply(subject=f"Re: {original_msg.subject}", body=html_body)

        idempotency.mark_success(idempotency_key)
        import json
        return json.dumps({"success": True, "action": "EmailReplied", "idempotencyKey": idempotency_key, "messageId": message_id})
    except Exception as e:
        idempotency.mark_failed(idempotency_key)
        raise e


@mcp.tool()
def forward_email(
    message_id: str, 
    to_recipients: str, 
    idempotency_key: str, 
    body_prefix: str = "", 
    cc_recipients: str = "", 
    use_signature: bool = True
) -> str:
    """Forward an email. to/cc recipients should be comma-separated strings."""
    idempotency.lock(idempotency_key)
    try:
        account = get_ews_client()
        original_msg = account.root.get(id=message_id)
        html_body = build_email_body(body_prefix, use_signature) if body_prefix else None
        
        to_list = [r.strip() for r in to_recipients.split(",") if r.strip()]
        cc_list = [r.strip() for r in cc_recipients.split(",") if r.strip()]
        
        original_msg.forward(
            subject=f"Fw: {original_msg.subject}",
            body=html_body,
            to_recipients=[Mailbox(email_address=addr) for addr in to_list],
            cc_recipients=[Mailbox(email_address=addr) for addr in cc_list]
        )

        idempotency.mark_success(idempotency_key)
        import json
        return json.dumps({"success": True, "action": "EmailForwarded", "idempotencyKey": idempotency_key, "messageId": message_id})
    except Exception as e:
        idempotency.mark_failed(idempotency_key)
        raise e


# ---------------------------------------------------------
# Management Operations
# ---------------------------------------------------------

@mcp.tool()
def mark_as_read(message_id: str, is_read: bool = True) -> str:
    """Mark an email as read or unread."""
    account = get_ews_client()
    try:
        item = account.root.get(id=message_id)
        item.is_read = is_read
        item.save(update_fields=['is_read'])
        import json
        return json.dumps({"success": True, "action": "MarkedAsRead" if is_read else "MarkedAsUnread"})
    except Exception as e:
        raise e

@mcp.tool()
def move_message(message_id: str, destination_folder: str) -> str:
    """Move an email to a destination folder (e.g. 'deleteditems', 'inbox')."""
    account = get_ews_client()
    try:
        item = account.root.get(id=message_id)
        dest = get_folder_by_name(account, destination_folder)
        item.move(dest)
        import json
        return json.dumps({"success": True, "action": "MessageMoved", "destination": dest.name})
    except Exception as e:
        raise e

@mcp.tool()
def delete_message(message_id: str, hard_delete: bool = False) -> str:
    """Delete an email. Default moves to 'Deleted Items' folder. Set hard_delete=True to permanently remove it."""
    account = get_ews_client()
    try:
        item = account.root.get(id=message_id)
        if hard_delete:
            item.delete()
        else:
            item.move_to_trash()
        import json
        return json.dumps({"success": True, "action": "MessageDeleted", "hard_delete": hard_delete})
    except Exception as e:
        raise e


@mcp.tool()
def batch_mark_as_read(message_ids: str, is_read: bool = True) -> str:
    """Batch mark multiple emails as read or unread. message_ids is comma-separated."""
    account = get_ews_client()
    ids = [i.strip() for i in message_ids.split(",") if i.strip()]
    if not ids:
        return '{"success": false, "error": "No message IDs provided."}'
    
    try:
        items = list(account.root.filter(id__in=ids))
        if not items:
            return '{"success": false, "error": "No matching messages found."}'
        
        for item in items:
            item.is_read = is_read
        
        account.bulk_update(items=[(item, ('is_read',)) for item in items])
        import json
        return json.dumps({"success": True, "count": len(items), "action": "BatchMarkRead" if is_read else "BatchMarkUnread"})
    except Exception as e:
        import json
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)


@mcp.tool()
def batch_move_messages(message_ids: str, destination_folder: str) -> str:
    """Batch move multiple emails to a destination folder. message_ids is comma-separated."""
    account = get_ews_client()
    ids = [i.strip() for i in message_ids.split(",") if i.strip()]
    if not ids:
        return '{"success": false, "error": "No message IDs provided."}'
    
    try:
        dest = get_folder_by_name(account, destination_folder)
        items = list(account.root.filter(id__in=ids))
        if not items:
            return '{"success": false, "error": "No matching messages found."}'
        
        account.bulk_move(items=items, to_folder=dest)
        import json
        return json.dumps({"success": True, "count": len(items), "destination": dest.name})
    except Exception as e:
        import json
        return json.dumps({"success": False, "error": str(e)}, ensure_ascii=False)


def serve_stdio():
    """Run FastMCP via stdio"""
    mcp.run(transport="stdio")
    
def serve_sse(port: int = 3100):
    """Run FastMCP via SSE (uvicorn internal)"""
    mcp.settings.port = port
    mcp.run(transport="sse")
